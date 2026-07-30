from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PURPLE = "4B1535"
PURPLE_DARK = "261019"
GOLD = "B78A2E"
GOLD_LIGHT = "E8D7A8"
CREAM = "F7F2E7"
WHITE = "FFFFFF"
INK = "272229"
GRAY = "6A646A"
GREEN = "37684D"
RED = "9C3D47"

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentos_oficiais"
OUT.mkdir(parents=True, exist_ok=True)


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text_color(cell, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    elif not value and keep is not None:
        p_pr.remove(keep)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("NorteP Pesquisa  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_document(doc: Document, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (
        ("Title", 30, PURPLE),
        ("Subtitle", 14, GOLD),
        ("Heading 1", 19, PURPLE),
        ("Heading 2", 14, PURPLE),
        ("Heading 3", 11.5, GOLD),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = f"NORTEP  •  {subtitle.upper()}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.name = "Aptos"
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor.from_string(PURPLE)
    bottom = OxmlElement("w:pBdr")
    edge = OxmlElement("w:bottom")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "8")
    edge.set(qn("w:color"), GOLD)
    bottom.append(edge)
    hp._p.get_or_add_pPr().append(bottom)
    add_page_number(section.footer.paragraphs[0])


def cover(doc: Document, title: str, subtitle: str, edition: str = "Versão 2026") -> None:
    for _ in range(3):
        doc.add_paragraph()
    logo = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(logo.rows[0])
    logo.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo.autofit = False
    logo.columns[0].width = Cm(4.2)
    cell = logo.cell(0, 0)
    cell.width = Cm(4.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, PURPLE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("N P")
    r.font.name = "Aptos Display"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
    doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(edition)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph()
    callout(doc, "Uso interno", "Material de formação e operação da equipe NorteP. Não contém senhas, chaves ou dados pessoais de entrevistados.", "gold")
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def callout(doc: Document, title: str, body: str, kind: str = "purple") -> None:
    colors = {
        "purple": (PURPLE, WHITE),
        "gold": (GOLD_LIGHT, PURPLE_DARK),
        "green": ("DCEBDD", GREEN),
        "red": ("F5DEE1", RED),
    }
    fill, text = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(text)
    r = p.add_run(body)
    r.font.color.rgb = RGBColor.from_string(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        shade(cell, PURPLE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        if widths:
            cell.width = Inches(widths[idx])
    for r_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            set_cell_margins(cell)
            if r_idx % 2:
                shade(cell, CREAM)
            p = cell.paragraphs[0]
            p.add_run(value)
            if widths:
                cell.width = Inches(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run("☐  ")
        r.font.color.rgb = RGBColor.from_string(GOLD)
        r.font.bold = True
        p.add_run(item)


def add_flow(doc: Document, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(steps))
    set_repeat_table_header(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, (title, body) in enumerate(steps):
        cell = table.cell(0, idx)
        shade(cell, PURPLE if idx % 2 == 0 else GOLD)
        set_cell_margins(cell, 180, 120, 180, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE if idx % 2 == 0 else PURPLE_DARK)
        r = p.add_run(body)
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(WHITE if idx % 2 == 0 else PURPLE_DARK)
    doc.add_paragraph()


def new_doc(subtitle: str) -> Document:
    doc = Document()
    configure_document(doc, subtitle)
    core = doc.core_properties
    core.author = "NorteP Pesquisa"
    core.company = "NorteP"
    core.subject = subtitle
    core.keywords = "NorteP, pesquisa de campo, LGPD, território, operação"
    return doc


def build_field_manual() -> Path:
    doc = new_doc("Formação e Operação de Campo")
    cover(
        doc,
        "Manual de Formação e Operação de Campo",
        "Método M.I.S.T. reformulado para pesquisa, escuta territorial e mobilização ética",
    )

    add_heading(doc, "1. Finalidade do manual")
    doc.add_paragraph(
        "Este manual transforma os materiais de agente mobilizador, supervisor e M.I.S.T. em um protocolo único para a NorteP. "
        "Ele orienta a equipe a ouvir pessoas, registrar dados confiáveis, proteger a autonomia do entrevistado e transformar a coleta em informação útil."
    )
    callout(
        doc,
        "Princípio central",
        "A confiança vale mais do que uma resposta. Ninguém deve ser pressionado, enganado, constrangido ou levado a acreditar que um benefício depende da participação.",
        "gold",
    )

    add_heading(doc, "2. O M.I.S.T. na operação NorteP")
    add_table(
        doc,
        ["Etapa", "O que significa", "Entrega prática"],
        [
            ["M — Mapear", "Definir território, público, horários e metas sem classificar pessoas por vulnerabilidade.", "Plano de campo e cobertura."],
            ["I — Investigar", "Ouvir com neutralidade, registrar respostas e ocorrências sem completar pelo entrevistado.", "Entrevistas e relatos íntegros."],
            ["S — Sistematizar", "Conferir sincronização, duração, recusas, qualidade e consentimentos.", "Banco de dados auditável."],
            ["T — Transformar", "Produzir diagnóstico, prioridades e aprendizados; separar fato, interpretação e recomendação.", "Relatórios de gestão e pesquisa."],
        ],
        [1.1, 2.2, 2.6],
    )
    doc.add_paragraph(
        "Os conceitos úteis dos documentos originais — escuta ativa, leitura territorial, treinamento prático, supervisão e consolidação diária — foram preservados. "
        "Expressões de confronto, “guerra”, manipulação, gatilhos ocultos ou exploração de vulnerabilidade foram substituídas por linguagem profissional, verificável e ética."
    )

    add_heading(doc, "3. Papéis e responsabilidades")
    add_table(
        doc,
        ["Papel", "Responsabilidade principal", "Limite"],
        [
            ["Administradora fundadora", "Governança do sistema, criação de administradores, visão integral, auditoria e cofre.", "Acesso exclusivo e intransferível."],
            ["Administrador secundário", "Pesquisas, coordenadores, supervisores, pesquisadores, territórios e relatórios.", "Não cria outro administrador."],
            ["Coordenador", "Gerencia sua cidade/região, supervisores, pesquisadores, pesquisas e desempenho da própria estrutura.", "Não vê equipes alheias."],
            ["Supervisor", "Acompanha pesquisadores vinculados, treina, corrige qualidade e registra ocorrências.", "Não cria coordenadores."],
            ["Pesquisador", "Aplica apenas pesquisas liberadas, registra respostas, recusas e sincronização.", "Não acessa painel administrativo."],
            ["Observador", "Acompanha indicadores agregados autorizados.", "Não altera dados nem acessos."],
        ],
        [1.35, 2.75, 1.65],
    )

    add_heading(doc, "4. Pesquisa, mobilização e relacionamento não são a mesma coisa")
    add_table(
        doc,
        ["Atividade", "Abertura correta", "Regra"],
        [
            ["Pesquisa quantitativa/qualitativa", "“Estamos ouvindo pessoas sobre o bairro e os serviços. Participar é voluntário.”", "Não induzir; registrar inclusive respostas desfavoráveis."],
            ["Pesquisa eleitoral", "Informar objetivo, voluntariedade e tratamento agregado de opinião política.", "Revisão jurídica e metodológica antes da ativação."],
            ["Formulário de relacionamento", "Informar que o link veio de apoiador/liderança e que contatos são opcionais.", "Cada uso de contato exige autorização separada."],
            ["Mobilização", "Dizer claramente quem promove a ação e que não há contratação ou benefício prometido.", "Convite pode ser recusado sem consequência."],
        ],
        [1.3, 2.5, 2.0],
    )

    add_heading(doc, "5. Proteção de dados e consentimento")
    add_bullets(
        doc,
        [
            "O entrevistado participa sem login, e-mail ou nome.",
            "Contato só é solicitado no final e apenas se a pessoa escolher receber resultados, conteúdos ou convites.",
            "Nome, WhatsApp e e-mail ficam separados dos resultados comuns, dentro do cofre autorizado.",
            "Opinião política é dado pessoal sensível: minimize a coleta, limite o acesso e analise de forma agrupada.",
            "Não fotografe documento, tela ou pessoa; não copie contatos do celular; não compartilhe planilhas por WhatsApp.",
            "Localização informada e local da coleta devem ser campos distintos. Geolocalização do aparelho depende de autorização específica.",
            "A pessoa pode interromper a entrevista e pedir retirada de contato pelo canal institucional.",
        ],
    )
    callout(
        doc,
        "Canal institucional",
        "Dúvidas ou pedidos sobre dados: pesquisadecamponortep@gmail.com. A coordenação deve registrar e encaminhar solicitações de acesso, correção ou retirada.",
        "purple",
    )

    add_heading(doc, "6. Preparação antes de sair")
    add_checklist(
        doc,
        [
            "Celular carregado, navegador atualizado e aplicativo instalado.",
            "Login testado com internet; pesquisa correta liberada e aberta uma vez para cache.",
            "Tema claro/noturno e tamanho da fonte ajustados.",
            "Território, horário, meta e ponto de apoio confirmados.",
            "Crachá e identificação disponíveis; roupa adequada e postura respeitosa.",
            "Rota segura, contatos do supervisor e protocolo de emergência conhecidos.",
            "Nenhum dado pessoal armazenado fora do NorteP.",
        ],
    )

    add_heading(doc, "7. Abertura humana da entrevista")
    add_flow(
        doc,
        [
            ("1. Cumprimente", "Tom natural, distância respeitosa."),
            ("2. Identifique-se", "Nome e NorteP Pesquisa."),
            ("3. Explique", "Tema, duração e uso agrupado."),
            ("4. Peça consentimento", "Aceite sim ou não sem insistir."),
        ],
    )
    callout(
        doc,
        "Roteiro sugerido",
        "“Bom dia/boa tarde. Meu nome é [nome] e faço parte da NorteP Pesquisa. Estamos ouvindo pessoas sobre [tema]. Leva cerca de [tempo] minutos. Participar é voluntário, você pode não responder ou parar quando quiser. Posso começar?”",
        "green",
    )
    add_heading(doc, "8. Como criar vínculo sem manipular", 2)
    add_bullets(
        doc,
        [
            "Ouça até o fim antes de tocar na tela.",
            "Use o nome somente quando a pessoa o tiver informado voluntariamente.",
            "Confirme entendimento: “Entendi corretamente que sua prioridade é…?”",
            "Ajuste o ritmo da fala, mas não imite gestos, sotaque ou emoção.",
            "Não invente prova social, autoridade, urgência, escassez ou promessa.",
            "Não corrija opinião política nem tente “ganhar” a conversa.",
            "Se a resposta for crítica, agradeça e registre com o mesmo cuidado.",
        ],
    )

    add_heading(doc, "9. Aplicação correta do questionário")
    add_steps(
        doc,
        [
            "Leia a pergunta como está escrita. Só explique o termo quando a instrução permitir.",
            "Em pergunta espontânea, não leia nomes ou alternativas.",
            "Em escolha única, marque uma resposta; em múltipla escolha, respeite o limite exibido.",
            "Use “prefere não responder”, “não sabe” ou equivalente; nunca invente resposta para avançar.",
            "Em texto aberto, registre a ideia completa. Uma letra isolada, texto sem sentido ou comentário do pesquisador não substitui a resposta.",
            "Use “manifestação espontânea” somente para algo dito sem pergunta e sem interpretação.",
            "Revise respostas obrigatórias antes de concluir.",
        ],
    )
    callout(
        doc,
        "Integridade da entrevista",
        "A NorteP pode sinalizar duração incompatível, repetição de respostas, texto sem sentido, volume incomum ou sequência improvável. O alerta pede revisão; não é condenação automática do pesquisador.",
        "gold",
    )

    add_heading(doc, "10. Recusa, interrupção e ocorrência")
    add_table(
        doc,
        ["Situação", "Ação no aplicativo", "Conduta"],
        [
            ["Não quer participar", "Registrar recusa e motivo quando informado.", "Agradecer; não insistir."],
            ["Precisa parar", "Registrar entrevista interrompida.", "Preservar o rascunho quando apropriado."],
            ["Erro técnico", "Salvar rascunho; registrar ocorrência.", "Não repetir respostas de memória."],
            ["Agressão, ameaça ou violência", "Registrar ocorrência com palavras-chave; alerta sincroniza quando houver sinal.", "Sair do local e priorizar segurança."],
            ["Pedido de retirada de dados", "Encaminhar ao canal institucional.", "Não prometer prazo que não controla."],
        ],
        [1.25, 2.1, 2.4],
    )

    add_heading(doc, "11. Trabalho sem internet")
    add_steps(
        doc,
        [
            "Antes de ir a campo, entre com internet e abra a pesquisa liberada.",
            "Durante a coleta, acompanhe o aviso “salva no aparelho”.",
            "Se fechar a página, use “continuar” para retomar o rascunho.",
            "Não limpe dados do navegador, não desinstale o app e não troque de navegador com pendências.",
            "Ao recuperar sinal, toque em sincronizar e aguarde o contador chegar a zero.",
            "Comunique ao supervisor qualquer pendência que permaneça após nova tentativa.",
        ],
    )

    add_heading(doc, "12. Supervisão e desenvolvimento da equipe")
    doc.add_paragraph(
        "A supervisão é apoio técnico, não competição. O ranking fica limitado à área da equipe e deve orientar treinamento, cobertura e reconhecimento — nunca exposição ou humilhação."
    )
    add_table(
        doc,
        ["Momento", "Prática", "Pergunta de feedback"],
        [
            ["Antes do turno", "Roleplay de 10 minutos e checagem do app.", "“Qual parte do roteiro ainda está insegura?”"],
            ["Primeira abordagem", "Eu faço: supervisor demonstra.", "“O que você percebeu na abertura?”"],
            ["Segunda abordagem", "Nós fazemos: apoio próximo.", "“Onde posso ajudar sem assumir sua fala?”"],
            ["Terceira abordagem", "Você faz: observação discreta.", "“O que funcionou e o que ajustaremos?”"],
            ["Fechamento", "Revisar volume, recusas, pendências e qualidade.", "“Qual aprendizado deve orientar amanhã?”"],
        ],
        [1.1, 2.2, 2.45],
    )

    add_heading(doc, "13. Indicadores que fazem sentido")
    add_table(
        doc,
        ["Indicador", "Cálculo/uso", "Cuidado"],
        [
            ["Entrevistas concluídas", "Volume por pessoa, pesquisa e território.", "Não mede qualidade sozinho."],
            ["Taxa de adesão", "Concluídas ÷ abordagens registradas.", "Registrar recusas é obrigatório para o cálculo."],
            ["Tempo mediano", "Mediana da duração das entrevistas.", "Muito rápido ou muito lento pede revisão."],
            ["Pendências offline", "Itens ainda no aparelho.", "Zerar antes do encerramento do turno."],
            ["Qualidade textual", "Respostas completas e coerentes.", "Não penalizar linguagem simples ou opinião divergente."],
            ["Cobertura territorial", "Entrevistas por cidade, região e bairro.", "Evitar concentração por conveniência."],
            ["Ocorrências", "Recusas, interrupções e alertas de segurança.", "Segurança prevalece sobre meta."],
        ],
        [1.25, 2.25, 2.25],
    )

    add_heading(doc, "14. Uso científico e relatório do mestrado")
    add_bullets(
        doc,
        [
            "Separar operação política, pesquisa acadêmica e relacionamento em bases/finalidades documentadas.",
            "Usar dados pessoais apenas quando houver base legal e consentimento específico; no relatório científico, trabalhar com dados anonimizados.",
            "Registrar instrumento, versão, período, territórios, treinamento, critérios de inclusão e perdas.",
            "Não inferir causalidade a partir de comparação simples entre desempenho e intenção de voto.",
            "Não direcionar intervenção a pessoas classificadas como vulneráveis. A pesquisa pode estudar contexto e boas práticas, não explorar fragilidades.",
            "Submeter o projeto ao comitê de ética quando aplicável e revisar regras eleitorais antes de ativar pesquisas eleitorais.",
            "Documentar limitações, recusas, viés de seleção, efeito do entrevistador e mudanças do questionário.",
        ],
    )

    add_heading(doc, "15. Avaliação prática do pesquisador")
    add_table(
        doc,
        ["Critério", "Peso", "Evidência"],
        [
            ["Consentimento e respeito", "25%", "Abertura clara, aceita recusa, não pressiona."],
            ["Fidelidade ao questionário", "25%", "Lê corretamente e não induz."],
            ["Qualidade do registro", "20%", "Texto coerente, local correto, ocorrências registradas."],
            ["Operação do aplicativo", "15%", "Rascunho, offline e sincronização usados corretamente."],
            ["Postura e segurança", "15%", "Identificação, autocontrole e acionamento do supervisor."],
        ],
        [2.0, 0.75, 3.0],
    )
    callout(
        doc,
        "Critério de aprovação sugerido",
        "Roleplay aprovado + entrevista acompanhada + nota mínima de 80%. Em falha de consentimento, falsificação ou segurança, interromper a coleta e reciclar antes do retorno.",
        "red",
    )

    add_heading(doc, "16. Cartão de bolso")
    add_checklist(
        doc,
        [
            "Eu me identifiquei e expliquei a finalidade?",
            "A pessoa aceitou livremente?",
            "Li sem induzir?",
            "Registrei local da coleta e residência separadamente?",
            "Usei “não sabe/prefere não responder” quando necessário?",
            "Registrei recusa, interrupção ou ocorrência?",
            "O rascunho está salvo e a fila está sincronizada?",
            "Nenhum dado pessoal ficou fora do cofre NorteP?",
        ],
    )

    path = OUT / "Manual_de_Formacao_e_Operacao_de_Campo_NorteP_2026.docx"
    doc.save(path)
    return path


def build_app_manual() -> Path:
    doc = new_doc("Manual de Uso do Aplicativo")
    cover(
        doc,
        "Manual de Uso do Aplicativo NorteP",
        "Acesso, pesquisas, territórios, coleta offline, mobilização, cofre e relatórios",
    )

    add_heading(doc, "1. Visão geral")
    doc.add_paragraph(
        "O NorteP é uma PWA: funciona no navegador e pode ser instalada no celular. A administração controla acessos, pesquisas e territórios; o entrevistado não cria conta. "
        "O acesso principal continua exclusivo da administradora fundadora."
    )
    add_flow(
        doc,
        [
            ("Convite", "Link correto para o papel."),
            ("Cadastro", "Nome, e-mail e senha."),
            ("Confirmação", "Abrir o e-mail recebido."),
            ("Aprovação", "Gestor ativa e vincula."),
            ("Uso", "Entrar sempre pelo mesmo canal."),
        ],
    )

    add_heading(doc, "2. Endereços de entrada")
    add_table(
        doc,
        ["Perfil", "Canal", "Quem envia"],
        [
            ["Administradora fundadora", "?acesso=principal", "Uso pessoal; não compartilhar."],
            ["Administrador secundário", "?acesso=administracao", "Somente a fundadora gera o convite."],
            ["Coordenador", "?acesso=coordenacao", "Fundadora ou administrador secundário."],
            ["Supervisor", "?acesso=supervisao", "Administrador ou coordenador responsável."],
            ["Observador", "?acesso=observador", "Administração."],
            ["Pesquisador", "?acesso=pesquisador", "Administração, coordenador ou supervisor responsável."],
            ["Eleitor/participante", "?mobilizacao=CÓDIGO", "Link público gerado para apoiador/liderança; sem login."],
        ],
        [1.55, 1.65, 2.5],
    )
    callout(
        doc,
        "Segurança do link",
        "O link escolhe a porta de entrada, mas não concede poder sozinho. O e-mail confirmado, o convite, o papel e as políticas do banco precisam concordar.",
        "gold",
    )

    add_heading(doc, "3. Cadastro e primeiro acesso")
    add_steps(
        doc,
        [
            "Abra o convite destinado ao seu e-mail.",
            "Toque em “Cadastrar” e informe nome completo, e-mail correto e uma senha forte.",
            "Abra a mensagem de confirmação. Verifique spam/lixo eletrônico se necessário.",
            "Ao voltar ao NorteP, aguarde a aprovação do responsável.",
            "Depois de aprovado, saia e entre novamente pelo canal do seu perfil.",
            "Use o olho da senha para conferir a digitação. Não compartilhe senha nem código.",
        ],
    )
    callout(
        doc,
        "Se a confirmação volta ao app, mas não libera",
        "Confira se o link pertence ao mesmo e-mail cadastrado, se não é uma mensagem antiga e se a aprovação foi feita. Em Yahoo ou filtros rígidos, aguarde alguns minutos e abra a mensagem mais recente.",
        "purple",
    )

    add_heading(doc, "4. O que cada perfil enxerga")
    add_table(
        doc,
        ["Perfil", "Pode", "Não pode"],
        [
            ["Fundadora", "Todas as entradas, visões, acessos, auditoria, cofre e criação de administradores.", "Não deve compartilhar a conta."],
            ["Administrador", "Criar coordenadores, supervisores e pesquisadores; editar pesquisas; ver operação geral.", "Criar outro administrador."],
            ["Coordenador", "Ver sua área; convidar supervisor/pesquisador; editar pesquisa permitida; acompanhar sua equipe.", "Ver equipe de outro coordenador."],
            ["Supervisor", "Ver sua equipe e área; convidar pesquisador; acompanhar qualidade e ocorrências.", "Criar coordenador ou supervisor."],
            ["Observador", "Ver painel agregado autorizado.", "Alterar pessoas, pesquisas ou contatos."],
            ["Pesquisador", "Ver pesquisas explicitamente liberadas; coletar e sincronizar.", "Acessar painel administrativo."],
        ],
        [1.25, 2.75, 1.85],
    )

    add_heading(doc, "5. Administração de acessos e territórios")
    add_steps(
        doc,
        [
            "Abra “Acessos e cadastros”.",
            "Escolha o papel e o responsável direto.",
            "Defina cidade, região e/ou bairro. Um perfil pode receber mais de um território.",
            "Gere o convite e envie somente ao e-mail indicado.",
            "Após confirmação, confira papel, vínculo, território e status ativo.",
            "Use suspender para interromper temporariamente; apagar acesso encerra a conta sem apagar o histórico já autorizado.",
        ],
    )
    add_table(
        doc,
        ["Vínculo", "Regra"],
        [
            ["Coordenador → supervisor", "Supervisor herda somente a área atribuída e vê sua equipe."],
            ["Coordenador → pesquisador", "Permitido quando o pesquisador responde diretamente ao coordenador."],
            ["Supervisor → pesquisador", "Permitido somente na própria estrutura."],
            ["Administração → coordenador", "Cidade/região/bairro são definidos no convite ou depois."],
            ["Fundadora → administrador", "Único caminho para novo administrador secundário."],
        ],
        [2.0, 3.8],
    )

    add_heading(doc, "6. Criar, editar, pausar e liberar pesquisas")
    add_steps(
        doc,
        [
            "Em “Pesquisas”, crie ou abra o questionário.",
            "Defina título, descrição, tipo, duração, texto de consentimento e território.",
            "Crie seções e perguntas; marque obrigatórias e configure condições.",
            "Salve como rascunho e teste no celular.",
            "Use “Liberar” para escolher pesquisadores específicos — pesquisa ativa não vai automaticamente para todos.",
            "Ative/pilote somente após revisão. Pausar retira da coleta sem apagar.",
            "Apagar é restrito à administração. Dados de teste devem ser separados e limpos antes do uso oficial.",
        ],
    )
    add_table(
        doc,
        ["Tipo", "Uso recomendado"],
        [
            ["Texto curto/longo", "Cidade, bairro, resposta espontânea, justificativa."],
            ["Sim/Não", "Filtro, consentimento ou condição."],
            ["Escolha única", "Uma resposta; segundo toque desmarca quando permitido."],
            ["Múltipla escolha", "Várias respostas, respeitando limite configurado."],
            ["Escala 0–10 / avaliação", "Intensidade ou qualidade."],
            ["Região/bairro", "Território informado."],
            ["Observação interna", "Nota do pesquisador que não deve ser lida."],
        ],
        [1.7, 4.1],
    )

    add_heading(doc, "7. Coleta pelo pesquisador")
    add_steps(
        doc,
        [
            "Abra a área do pesquisador e toque em atualizar.",
            "Escolha uma das pesquisas liberadas.",
            "Leia o consentimento e registre a resposta.",
            "Responda a etapa atual; ao avançar, a tela sobe automaticamente para o início da próxima.",
            "Perguntas obrigatórias precisam de resposta válida; use “não sabe/prefere não responder” quando disponível.",
            "Ao concluir, confira o código ENT-AAAA-000001, a duração e o status de sincronização.",
        ],
    )
    callout(
        doc,
        "Contato opcional",
        "Nome, WhatsApp ou e-mail só aparecem se o entrevistado quiser receber resultados/conteúdos e autorizar o armazenamento. A pesquisa continua anônima quando a resposta for “Não”.",
        "green",
    )

    add_heading(doc, "8. Rascunho, queda da página e modo offline")
    add_table(
        doc,
        ["Situação", "O que o NorteP faz", "O que a pessoa deve fazer"],
        [
            ["Sem internet", "Mantém formulário e salva entrevista no aparelho.", "Continuar e não limpar o navegador."],
            ["Página fechou", "Oferece continuar ou recomeçar.", "Continuar para preservar respostas."],
            ["Voltou o sinal", "Tenta sincronizar pendências.", "Tocar em sincronizar e aguardar zero."],
            ["Pesquisa mudou", "Nova versão chega quando houver sinal.", "Atualizar antes do turno; não misturar versões."],
            ["Erro persistente", "Mantém item na fila.", "Anotar código/aparelho e avisar supervisor."],
        ],
        [1.25, 2.3, 2.25],
    )
    add_checklist(
        doc,
        [
            "Entrar uma vez com internet.",
            "Abrir a pesquisa antes de perder o sinal.",
            "Não usar aba anônima/privada.",
            "Não limpar cache/dados com pendências.",
            "Sincronizar no mesmo aparelho e navegador.",
        ],
    )

    add_heading(doc, "9. Recusas, interrupções e alertas")
    doc.add_paragraph(
        "Abordagem, recusa e interrupção fazem parte da operação. Registrar esses eventos permite calcular adesão real e entender falhas sem transformar a meta em pressão indevida."
    )
    add_bullets(
        doc,
        [
            "Recusa: registre quando a pessoa não inicia.",
            "Interrompida: registre quando a entrevista começou e não terminou.",
            "Ocorrência: descreva problema técnico, conflito ou risco.",
            "Alerta de segurança: palavras relacionadas a agressão, ameaça ou violência destacam o evento para a gestão.",
            "Sem internet: o alerta fica no aparelho e chega quando sincronizar.",
        ],
    )

    add_heading(doc, "10. Painéis, território e CSV")
    add_bullets(
        doc,
        [
            "Visão geral: volume, tempo, recusas, interrupções e alertas.",
            "Cobertura territorial: ranking por cidade, região e bairro — sem mapa físico nesta versão.",
            "Equipe: entrevistas e tempo por pesquisador, limitado ao escopo do gestor.",
            "Resultados: percentuais de perguntas fechadas e respostas abertas preservadas.",
            "Exportar CSV: baixa uma tabela compatível com Excel, Power BI, Python e Looker Studio.",
            "O CSV operacional não deve incluir contatos do cofre.",
        ],
    )

    add_heading(doc, "11. Mobilização por apoiadores e lideranças")
    add_steps(
        doc,
        [
            "A administração abre “Mobilização”.",
            "Cadastra apoiador/liderança e território; vídeo de agradecimento é opcional.",
            "O NorteP cria um link individual para enviar por WhatsApp.",
            "O participante abre sem cadastro e responde o formulário público.",
            "Autorizações de conteúdo, reuniões, voluntariado e uso acadêmico são separadas.",
            "O painel contabiliza respostas e adesões por parceiro e território.",
        ],
    )
    callout(
        doc,
        "Transparência obrigatória",
        "O formulário informa que participação/voluntariado não é contratação nem promessa de pagamento. Compartilhar conteúdo ou apoiar é sempre opcional.",
        "gold",
    )

    add_heading(doc, "12. Cofre de contatos")
    add_steps(
        doc,
        [
            "Abra “Cofre” com uma chave própria, diferente da senha do aplicativo.",
            "A fundadora libera responsáveis individualmente.",
            "Cada responsável cria sua própria chave; nenhuma chave deve ser compartilhada.",
            "A sessão do cofre expira e cada abertura/consulta fica registrada.",
            "Use os contatos somente para a finalidade autorizada.",
            "Revogue o acesso quando a função mudar ou o trabalho terminar.",
        ],
    )
    callout(
        doc,
        "Nunca fazer",
        "Não exportar o cofre para grupos, planilhas pessoais, WhatsApp ou e-mail sem controle. Não guardar a chave em mensagem, bloco de notas público ou junto da senha.",
        "red",
    )

    add_heading(doc, "13. Configurações e acessibilidade")
    add_bullets(
        doc,
        [
            "Configurações aparecem em todos os perfis.",
            "Tema claro e noturno ficam salvos no aparelho.",
            "Texto normal e A+ maior ajudam em telas pequenas.",
            "Botões, ícones e faixas mantêm contraste no modo noturno.",
            "O menu da conta reúne “Sair do aplicativo” e, quando permitido, “Descadastrar meu acesso”.",
        ],
    )

    add_heading(doc, "14. Solução rápida de problemas")
    add_table(
        doc,
        ["Problema", "Verificação"],
        [
            ["Não chegou confirmação", "Spam; endereço digitado; mensagem mais recente; aguardar o provedor."],
            ["Confirmou, mas espera aprovação", "Gestor precisa ativar e vincular o perfil correto."],
            ["Entrou como perfil errado", "Sair; usar o canal correto; conferir papel no painel."],
            ["Pesquisa não apareceu", "Confirmar liberação individual, status ativo/piloto e tocar em atualizar."],
            ["Teclado fecha", "Atualizar a página/app; testar navegador atualizado; registrar etapa e aparelho."],
            ["Não sincroniza", "Voltar à internet, manter o mesmo navegador e tocar em sincronizar."],
            ["Senha esquecida", "Usar recuperar senha e abrir o e-mail mais recente."],
            ["Tela sem contraste", "Alternar tema; atualizar; informar tela e botão afetado."],
        ],
        [2.0, 3.8],
    )

    add_heading(doc, "15. Checklist diário por função")
    add_heading(doc, "Pesquisador", 2)
    add_checklist(doc, ["Pesquisa correta", "Bateria e acesso", "Rascunhos revisados", "Pendências zeradas", "Ocorrências enviadas"])
    add_heading(doc, "Supervisor/coordenador", 2)
    add_checklist(doc, ["Equipe e território corretos", "Convites vinculados", "Alertas revisados", "Qualidade e duração verificadas", "Feedback realizado"])
    add_heading(doc, "Administração", 2)
    add_checklist(doc, ["Acessos por menor privilégio", "Pesquisas revisadas", "Exportações protegidas", "Cofre auditado", "Mudanças e incidentes documentados"])

    add_heading(doc, "16. Canal de suporte")
    doc.add_paragraph(
        "Ao relatar erro, informe: perfil, aparelho, navegador, horário, pesquisa, etapa, condição on-line/offline e mensagem exibida. "
        "Não envie senha, chave do cofre ou dados do entrevistado."
    )
    callout(doc, "Contato", "pesquisadecamponortep@gmail.com", "purple")

    path = OUT / "Manual_de_Uso_do_App_NorteP_2026.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    field = build_field_manual()
    app = build_app_manual()
    print(field)
    print(app)
